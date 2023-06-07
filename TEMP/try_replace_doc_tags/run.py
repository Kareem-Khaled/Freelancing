from docx import Document
import openpyxl

def replace_tags(file_path, tags_file_path):
    doc = Document(file_path)

    tags = {}
    workbook = openpyxl.load_workbook(tags_file_path)
    sheet = workbook.active

    for row in sheet.iter_rows(values_only=True):
        if len(row) >= 2:
            tag, value = row[:2]
            if tag:
                tags[tag] = value

    for tag, value in tags.items():
        # Replace tags in document body
        for paragraph in doc.paragraphs:
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, str(value))

        # Replace tags in headers and footers
        for section in doc.sections:
            header = section.header
            if header is not None:
                for paragraph in header.paragraphs:
                    for run in paragraph.runs:
	                    print(paragraph.runs.text)

					if tag in paragraph.text:
                        paragraph.text = paragraph.text.replace(tag, str(value))

            footer = section.footer
            if footer is not None:
                for paragraph in footer.paragraphs:
                    if tag in paragraph.text:
                        paragraph.text = paragraph.text.replace(tag, str(value))

    new_file_path = file_path.replace('.docx', '_new.docx')
    doc.save(new_file_path)
    print(f"New file saved as '{new_file_path}'.")

# Example usage:
file_path = 'document.docx'
tags_file_path = 'tags.xlsx'

replace_tags(file_path, tags_file_path)
